from docx import Document
import argparse
import json
import os
import glob 
import win32com.client
from docx2pdf import convert


# Set up argument parser
def parse_arguments():
    parser = argparse.ArgumentParser(description="Process a file path and a config JSON.")
    parser.add_argument(
        "--path",
        type=str,
        required=True,
        help="Path to the file or directory."
    )
    
    parser.add_argument(
        "--odir",
        type=str,
        default= "outputs",
        help="Path to the file or directory."
    )
    parser.add_argument(
        "--config",
        type=str,
        required=True,
        help="Path to the JSON configuration file or folder."
    )
    parser.add_argument(
        "--pdf",
        action='store_true',
        help="Wether to save as PDF or not."
    )
    return parser.parse_args()



def docx_to_pdf(input_file, output_file):
    convert(input_file, output_file)


def modify_doc(doc: Document, dict_values:dict, search_areas = ["paragraph","table","sections"]):
    # Load the DOCX file
    # doc = Document('example.docx')

    # Modify text in the document
    if "paragraph" in search_areas:
        for paragraph in doc.paragraphs:
            for k,v in dict_values.items():
                # print(k, paragraph.text, k in paragraph.text)
                if k in paragraph.text:
                    paragraph.text = paragraph.text.replace(k, v)
    if "table" in search_areas:
        # Search in tables
        for table_index, table in enumerate(doc.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for k,v in dict_values.items():
                        # print(k, paragraph.text, k in paragraph.text)
                        if k in cell.text:
                            cell.text = cell.text.replace(k, v)
                        
    if "section" in search_areas:
        # Search in headers and footers
        for section_index, section in enumerate(doc.sections):
            # Header
            for i, header in enumerate(section.header.paragraphs):
                for k,v in dict_values.items():
                    # print(k, paragraph.text, k in paragraph.text)
                    if k in header.text:
                        header.text = header.text.replace(k, v)
            # Footer
            for i, footer in enumerate(section.footer.paragraphs):
                for k,v in dict_values.items():
                    # print(k, paragraph.text, k in paragraph.text)
                    if k in footer.text:
                        footer.text = footer.text.replace(k, v)
    return doc


def parse_save_doc(input_path, json_config_path, odir, topdf):
    json_config = json.load(open(json_config_path))
    doc = Document(input_path)
    doc = modify_doc(doc=doc, dict_values=json_config)
    bname = os.path.basename(json_f).split(".json")[0]
    out_path = odir + "/" + input_path.replace(".docx",f"_compiled_{bname}.docx")
    doc.save(out_path)
    if topdf:
        # Example usage
        docx_to_pdf(out_path, out_path.replace(".docx",".pdf"))

if __name__=="__main__":
    args = parse_arguments()
    if not os.path.exists(args.odir):
        os.makedirs(args.odir)
    
    if os.path.isdir(args.config):
        for json_f in glob.glob(args.config + "/*.json"):
            parse_save_doc(args.path, json_f, args.odir, args.pdf)
    else:
        parse_save_doc(args.path, args.config, args.odir, args.pdf)
